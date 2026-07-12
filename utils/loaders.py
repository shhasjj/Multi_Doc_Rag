from langchain_community.document_loaders import PyPDFLoader
from docx import Document
from langchain_core.documents import Document as LCDocument
import os
import warnings

warnings.filterwarnings("ignore", category=DeprecationWarning)

def load_pdf(file_path):
    """
    Load a PDF file and return LangChain Document objects.
    """
    loader = PyPDFLoader(file_path)
    return loader.load()


def load_docx(file_path):
    """
    Load a DOCX file and convert it to LangChain Documents.
    """
    doc = Document(file_path)

    text = "\n".join(
        paragraph.text
        for paragraph in doc.paragraphs
        if paragraph.text.strip()
    )

    return [
        LCDocument(
            page_content=text,
            metadata={
                "source": os.path.basename(file_path),
                "page": 1
            }
        )
    ]


def load_documents(file_paths):
    """
    Load multiple PDF and DOCX files.
    """

    documents = []

    for file_path in file_paths:

        if file_path.endswith(".pdf"):
            documents.extend(load_pdf(file_path))

        elif file_path.endswith(".docx"):
            documents.extend(load_docx(file_path))

    return documents